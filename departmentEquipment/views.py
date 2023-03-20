from django.conf import settings # import the settings file
from django.contrib import messages
from django.shortcuts import render, redirect, get_object_or_404, HttpResponseRedirect
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, QueryDict, JsonResponse
from django.urls import reverse_lazy
from urllib.parse import urlencode
from django.views.generic import ListView, DetailView
from django.views.generic.edit import CreateView, UpdateView, DeleteView
from django.db.models import Q
from .models import Equipment, Status, Category
from employees.models import Employee
from .forms import EquipmentCreateForm, EquipmentUpdateForm
import os
import docx
from docx.enum.section import WD_ORIENT


@login_required(login_url="/")
def equipment_list(request):
    ADMIN_SITE_NAME = settings.DEFAULT_SITE_NAMING
    template = 'departmentEquipment/list.html'
    equipments = Equipment.objects.all()
    categories = Category.objects.all()
    statuses = Status.objects.all()
    employees = Employee.objects.all()
    sort_dict = (
        {'id':1, 'key':'name','value':'Назва'},
        {'id':2, 'key':'category__name','value':'Категорія'},
        {'id':3, 'key':'operationDate','value':'Дата вв. в експлуатацію'},
        {'id':4, 'key':'status__name','value':'Статус'},
        {'id':5, 'key':'location__name','value':'Місце закріплення'},
        {'id':6, 'key':'currentLocation__name','value':'Місцезнаходження'},
        {'id':7, 'key':'responsible__lastname','value':'Матеріально Відповідальний'},
        {'id':8, 'key':'fixed__lastname','value':'За ким закріплено'},
        {'id':9, 'key':'employee__lastname','value':'Користувач'},
        {'id':10, 'key':'inventoryNumber','value':'Інвентарний номер'},
        {'id':11, 'key':'internalNumber','value':'Внутрішній номер'},
        {'id':12, 'key':'serialNumber','value':'Заводський номер'},
        )

    # Отримати параметри запиту GET
    category_filters = request.GET.getlist('category[]')
    status_filters = request.GET.getlist('status[]')
    sort_field = request.GET.get('sort', 'name')
    sort_order = request.GET.get('order', 'asc')
    search_query = request.GET.get('q', '')
    
    # Фільтрувати дані за категорією
    if category_filters:
        equipments = equipments.filter(category__name__in=category_filters)
    
    # Фільтрувати дані за статусом
    if status_filters:
        equipments = equipments.filter(status__name__in=status_filters)

    # Фільтрувати дані за пошуком
    if search_query:
        equipments = equipments.filter(
                    Q(name__icontains=search_query) |
                    Q(location__name__icontains=search_query) |
                    Q(currentLocation__name__icontains=search_query) |
                    Q(responsible__lastname__icontains=search_query) |
                    Q(fixed__lastname__icontains=search_query) |
                    Q(employee__lastname__icontains=search_query) |
                    Q(inventoryNumber__icontains=search_query) |
                    Q(internalNumber__icontains=search_query) |
                    Q(serialNumber__icontains=search_query)
                    )
    
    # Сортувати дані за вибраним полем
    ordering = (sort_field, '-' + sort_field)[sort_order == 'desc']
    equipments = equipments.order_by(ordering)
    
    # Передати дані в шаблон
    context = {
        'ADMIN_SITE_NAME': ADMIN_SITE_NAME,
        'equipments': equipments,
        'categories': categories,
        'statuses': statuses,
        'employees': employees,
        'selected_categories':category_filters,
        'selected_statuses':status_filters,
        'sort_dict':sort_dict,
        'sort_field': sort_field,
        'sort_order': sort_order,
        'search_query': search_query
    }
    return render(request, template, context)


@login_required(login_url="/")
def equipment_detail(request, pk):
    ADMIN_SITE_NAME = settings.DEFAULT_SITE_NAMING
    template = 'departmentEquipment/detail.html'
    equipment = get_object_or_404(Equipment, id=pk)
    elements = Equipment.objects.all().filter(destinationEquipment=pk)
    history = equipment.history.all().order_by('history_date')

    context = {
        'ADMIN_SITE_NAME': ADMIN_SITE_NAME,
        'equipment': equipment,
        'elements':elements,
        'history': history
    }
    return render(request,template, context)


class EquipmentCreateView(CreateView):
    model = Equipment
    template_name = 'departmentEquipment/add.html'
    form_class = EquipmentCreateForm

    def get_context_data(self, **kwargs):
        context = super(EquipmentCreateView, self).get_context_data(**kwargs)
        return context

    def form_valid(self, form):
        if self.request.POST.get('_save'):
            messages.success(self.request, 'Дані було успішно збережено.')
        if self.request.POST.get('_dismiss'):
            messages.success(self.request, 'Ви відмінили запит на створення')
        return super(EquipmentCreateView, self).form_valid(form)

    def get_success_url(self):
        if self.request.POST.get('_save'):
            return reverse_lazy('departmentEquipment:list')
        if self.request.POST.get('_dismiss'):
            return reverse_lazy('departmentEquipment:list')

    def dispatch(self, request, *args, **kwargs):
        if request.user.is_authenticated == True:
            if request.method.lower() in self.http_method_names:
                handler = getattr(self, request.method.lower(), self.http_method_not_allowed)
            else:
                handler = self.http_method_not_allowed
            return handler(request, *args, **kwargs)
        else:
            return redirect("signin")


class EquipmentUpdateView(UpdateView):
    model = Equipment
    template_name = 'departmentEquipment/update.html'
    form_class = EquipmentUpdateForm

    @property
    def has_permission(self):
        return self.user.is_active

    def get_context_data(self, **kwargs):
        context = super(EquipmentUpdateView, self).get_context_data(**kwargs)
        return context

    def form_valid(self, form):
        equipment = self.get_object()
        if self.request.POST.get('_save'):
            messages.success(self.request, '\"{}\" було успішно змінено.'.format(equipment.name))
        if self.request.POST.get('_dismiss'):
            messages.success(self.request, 'Ви відмінили запит на зміну')
        return super(EquipmentUpdateView, self).form_valid(form)

    def get_success_url(self):
        if self.request.POST.get('_save'):
            return reverse_lazy('departmentEquipment:list')
        if self.request.POST.get('_dismiss'):
            return reverse_lazy('departmentEquipment:list')

    def dispatch(self, request, *args, **kwargs):
        if request.user.is_authenticated == True:
            if request.method.lower() in self.http_method_names:
                handler = getattr(self, request.method.lower(), self.http_method_not_allowed)
            else:
                handler = self.http_method_not_allowed
            return handler(request, *args, **kwargs)
        else:
            return redirect("signin")


class EquipmentDeleteView(DeleteView):
    model = Equipment
    template_name = 'departmentEquipment/confirm_delete.html'

    def get_success_url(self):
        return reverse_lazy('departmentEquipment:list')

    def post(self, request, *args, **kwargs):
        if self.request.POST.get('_cancel'):
            url = self.get_success_url()
            return HttpResponseRedirect(url)
        else:
            return super(EquipmentDeleteView, self).post(request, *args, **kwargs)

    def dispatch(self, request, *args, **kwargs):
        if request.user.is_authenticated == True:
            if request.method.lower() in self.http_method_names:
                handler = getattr(self, request.method.lower(), self.http_method_not_allowed)
            else:
                handler = self.http_method_not_allowed
            return handler(request, *args, **kwargs)
        else:
            return redirect("signin")



'''  History   '''
@login_required(login_url="/")
def equipment_history_list(request):
    ADMIN_SITE_NAME = settings.DEFAULT_SITE_NAMING
    template = 'departmentEquipment/history_list.html'
    history = Equipment.history.filter(history_type='-')

    context = {
        'history': history
    }
    return render(request, template, context)



def equipment_get_formular(request, pk):
    equipment = get_object_or_404(Equipment, id=pk)

    #formular = docx.Document('static/documents/FormularTemplate.docx')
    formularTemplate = docx.Document('static/documents/FormularTemplate.docx')
    formularTemplate.save('static/documents/Formular {}.docx'.format(equipment.name))

    formular = docx.Document('static/documents/Formular {}.docx'.format(equipment.name))

    




    section = formular.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_height = section.page_width
    section.page_width = section.page_height
    section.page_height = new_height

    table = formular.add_table(rows=1, cols=1)
    table.cell(0, 0).text = 'ФОРМУЛЯР \n {}'.format(equipment.name)

    formular.add_heading(equipment.name, 0)
    table = formular.add_table(rows=3, cols=3)
    table.cell(0, 0).text = 'Перший рядок, перша комірка'   
    table.cell(0, 1).text = 'Перший рядок, друга комірка'
    table.cell(0, 2).text = 'Перший рядок, третя комірка'
    table.cell(1, 0).text = 'Другий рядок, перша комірка'
    table.cell(1, 1).text = 'Другий рядок, друга комірка'
    table.cell(1, 2).text = 'Другий рядок, третя комірка'
    table.cell(2, 0).text = 'Третій рядок, перша комірка'
    table.cell(2, 1).text = 'Третій рядок, друга комірка'
    table.cell(2, 2).text = 'Третій рядок, третя комірка'

    #formular.add_paragraph(equipment.name)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=Formular \"{}\".docx'.format(equipment.name)
    formular.save(response)
    return response

    #history = Equipment.history.filter(history_type='-')
    #context = {
    #   'history': history
    #}

    #return render(request, template, context)


    #formular = docx.Document(template_path)

    #equipment = get_object_or_404(Equipment, id=pk)

    #template = Template('{{ var_name }}')

    #for paragraph in doc.paragraphs:
    #   if template.substitute(var_equipment_name='') in paragraph.text:
    #        paragraph.text = template.substitute(var_equipment_name=equipment.name)


    #doc.save('../static/documents/Formular\"{}\".doc'.format(equipment.name))


